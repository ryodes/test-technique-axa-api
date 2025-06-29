from flask import Flask, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime
import io

app = Flask(__name__)
CORS(app) 

# Config de la base SQLite
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///devis.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

class Devis(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    type_ouvrage = db.Column(db.String(50))
    nom_client = db.Column(db.String(100))
    numero_opportunite = db.Column(db.String(50))
    garantie = db.Column(db.String(50))
    destination_ouvrage = db.Column(db.String(100))
    type_travaux = db.Column(db.String(50))
    cout_ouvrage = db.Column(db.Float)
    presence_existant = db.Column(db.Boolean)
    client_vip = db.Column(db.Boolean)
    rcmo = db.Column(db.Boolean)
    tarif_trc = db.Column(db.Float)
    tarif_do = db.Column(db.Float)
    adresse_chantier = db.Column(db.String(255))
    description = db.Column(db.Text)

@app.route('/api/devis', methods=['GET'])
def get_devis():
    devis_liste = Devis.query.all()
    result = []

    for devis in devis_liste:
        result.append({
            "id": devis.id,
            "type_ouvrage": devis.type_ouvrage,
            "nom_client": devis.nom_client,
            "numero_opportunite": devis.numero_opportunite,
            "garantie": devis.garantie,
            "destination_ouvrage": devis.destination_ouvrage,
            "type_travaux": devis.type_travaux,
            "cout_ouvrage": devis.cout_ouvrage,
            "presence_existant": devis.presence_existant,
            "client_vip": devis.client_vip,
            "rcmo": devis.rcmo,
            "tarif_trc": devis.tarif_trc,
            "tarif_do": devis.tarif_do,
            "adresse_chantier": devis.adresse_chantier,
            "description": devis.description,
        })

    return jsonify(result)

@app.route("/api/devis", methods=["POST"])
def create_devis():
    data = request.json
    nouveau_devis = Devis(
        type_ouvrage=data.get('typeOuvrage'),
        nom_client=data.get('nomClient'),
        numero_opportunite=data.get('numeroOpportunite'),
        garantie=data.get('garantie'),
        destination_ouvrage=data.get('destinationOuvrage'),
        type_travaux=data.get('typeTravaux'),
        cout_ouvrage=data.get('coutOuvrage'),
        presence_existant=data.get('existant') == 'oui',
        client_vip=data.get('vip') == 'oui',
        rcmo=data.get('rcmo') == 'oui',
        tarif_trc=data.get('tarifTRC'),
        tarif_do=data.get('tarifDO'),
        adresse_chantier=data.get('adresseChantier'),
        description=data.get('description')
    )

    db.session.add(nouveau_devis)
    db.session.commit()

    return jsonify({"message": "Devis reçu", "data": data}), 200

@app.route('/api/devis/<int:devis_id>/pdf', methods=['GET'])
def generate_pdf(devis_id):
    devis = Devis.query.get_or_404(devis_id)

    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
    pdf.add_font("DejaVu", "B", "DejaVuSans-Bold.ttf", uni=True)
    pdf.add_font("DejaVu", "I", "DejaVuSerif-Italic.ttf", uni=True)
    pdf.set_font("DejaVu", size=10)
    
    # -------
    def add_section_title(title):
        pdf.set_font("DejaVu", style="B", size=10)
        pdf.cell(0, 10, txt=title, ln=True)
        pdf.set_font("DejaVu", size=10)

    pdf.image("logo.png", x=(210 - 20) / 2, y=10, w=20)
    pdf.ln(30)

    pdf.set_text_color(0, 0, 143)
    pdf.set_font("DejaVu", size=12)
    pdf.cell(0, 10, txt="TARIFICATION INDICATIVE", ln=True)

    pdf.set_text_color(0, 0, 0)
    pdf.set_font("DejaVu", size=10)
    pdf.multi_cell(0, 8, "Tarification indicative (*) sur la base d’un risque conforme aux paramètres suivants :")

    pdf.cell(0, 8, f"Type d’ouvrage : {devis.type_ouvrage}", ln=True)
    pdf.cell(0, 8, f"Types de travaux réalisés : {devis.type_travaux}", ln=True)
    pdf.cell(0, 8, f"Coût du chantier : {devis.cout_ouvrage} €", ln=True)
    pdf.cell(0, 8, f"Présence d’existant : {'Oui' if devis.presence_existant else 'Non'}", ln=True)
    pdf.cell(0, 8, f"Garantie choisie : {devis.garantie}", ln=True)
    pdf.cell(0, 8, f"Description de l’ouvrage : {devis.description}", ln=True)
    pdf.multi_cell(0, 8, f"Adresse du chantier : {devis.adresse_chantier}")

    # --- GARANTIES ---
    pdf.cell(0, 8, "Garanties Tous Risques chantier", ln=True)
    add_section_title("MONTANTS DE GARANTIES (exprimés en €)")
    pdf.cell(0, 8, "Dommages matériels à l'ouvrage : xxxxxxxxx", ln=True)
    pdf.cell(0, 8, "Responsabilité civile (tous dommages confondus) : xxxxxxxxx", ln=True)
    pdf.cell(0, 8, "Maintenance-visite : xxxxxxxxx", ln=True)
    pdf.cell(0, 8, "Mesure conservatoire : xxxxxxxxx", ln=True)

    pdf.ln(8)
    add_section_title("MONTANTS DE FRANCHISES (par sinistre exprimés en €)")
    pdf.cell(0, 8, "Dommages subis par les ouvrages de bâtiment : xxxxxxx", ln=True)
    pdf.cell(0, 8, "Catastrophes naturelles : montant défini par la loi", ln=True)
    pdf.cell(0, 8, "Responsabilité civile (1)", ln=True)
    pdf.cell(0, 8, "  - Assuré maître d'ouvrage : xxxxxxx", ln=True)
    pdf.cell(0, 8, "  - Assurés intervenants: SANS", ln=True)
    pdf.cell(0, 8, "Maintenance-visite : xxxxxxx", ln=True)
    pdf.ln(4)
    pdf.cell(0, 8, "     (1) Ces franchises s'appliquent pour des dommages autres que corporels", ln=True)

    pdf.ln(10)
    pdf.set_font("DejaVu", style="I", size=10)
    pdf.cell(0, 8, f"Date de simulation de tarif : le {datetime.today().strftime('%d/%m/%Y')}", ln=True)
    pdf.cell(0, 8, "(*) Cette tarification est faite sous réserve d’acceptation du risque par la compagnie", ln=True)
    # -------

    pdf_bytes = pdf.output(dest='S').encode('latin1')
    buffer = io.BytesIO(pdf_bytes)
    buffer.seek(0)

    return send_file(buffer, mimetype='application/pdf', as_attachment=True,
                     download_name=f"Proposition_commerciale_{devis.numero_opportunite}_{datetime.today().strftime('%d/%m/%Y:%Hh%M')}.pdf")

@app.route('/api/devis/<int:devis_id>/word')
def generate_word(devis_id):
    devis = Devis.query.get_or_404(devis_id)

    doc = Document()

    # Logo (facultatif, uniquement si Word accepte le PNG depuis le chemin local)
    try:
        doc.add_picture('logo.png', width=Inches(1.2))
    except Exception as e:
        print("Logo non chargé :", e)

    # Titre
    doc.add_paragraph("TARIFICATION INDICATIVE", style='Title')

    doc.add_paragraph(
        "Tarification indicative (*) sur la base d’un risque conforme aux paramètres suivants :",
        style='Normal'
    )

    doc.add_paragraph(f"Type d’ouvrage : {devis.type_ouvrage}")
    doc.add_paragraph(f"Types de travaux réalisés : {devis.type_travaux}")
    doc.add_paragraph(f"Coût du chantier : {devis.cout_ouvrage} €")
    doc.add_paragraph(f"Présence d’existant : {'Oui' if devis.presence_existant else 'Non'}")
    doc.add_paragraph(f"Garantie choisie : {devis.garantie}")
    doc.add_paragraph(f"Description de l’ouvrage : {devis.description}")
    doc.add_paragraph(f"Adresse du chantier : {devis.adresse_chantier}")

    # Section garanties
    doc.add_paragraph("Garanties Tous Risques chantier").bold = True
    doc.add_paragraph("MONTANTS DE GARANTIES (exprimés en €)").underline = True
    doc.add_paragraph("Dommages matériels à l'ouvrage : xxxxxxxxx")
    doc.add_paragraph("Responsabilité civile (tous dommages confondus) : xxxxxxxxx")
    doc.add_paragraph("Maintenance-visite : xxxxxxxxx")
    doc.add_paragraph("Mesure conservatoire : xxxxxxxxx")

    doc.add_paragraph("")

    doc.add_paragraph("MONTANTS DE FRANCHISES (par sinistre exprimés en €)").underline = True
    doc.add_paragraph("Dommages subis par les ouvrages de bâtiment : xxxxxxx")
    doc.add_paragraph("Catastrophes naturelles : montant défini par la loi")
    doc.add_paragraph("Responsabilité civile (1)")
    doc.add_paragraph("  - Assuré maître d'ouvrage : xxxxxxx")
    doc.add_paragraph("  - Assurés intervenants : SANS")
    doc.add_paragraph("Maintenance-visite : xxxxxxx")
    doc.add_paragraph("(1) Ces franchises s'appliquent pour des dommages autres que corporels")

    doc.add_paragraph("")
    doc.add_paragraph(f"Date de simulation de tarif : le {datetime.today().strftime('%d/%m/%Y')}")
    doc.add_paragraph("(*) Cette tarification est faite sous réserve d’acceptation du risque par la compagnie")

    # Envoi en mémoire
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True, 
                     download_name=f"Proposition_commerciale_{devis.numero_opportunite}_{datetime.today().strftime('%d/%m/%Y:%Hh%M')}.docx")

@app.route("/reset-devis", methods=["POST"])
def reset_devis():
    token = request.args.get("token")
    if token != "test123":
        return {"error": "Accès refusé"}, 403

    Devis.query.delete()
    db.session.commit()
    return {"message": "Base devis vidée ✅"}

if __name__ == "__main__":
    app.run(debug=True)

with app.app_context():
    db.create_all()
